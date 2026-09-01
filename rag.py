"""Core ingestion and retrieval logic for the claims RAG application."""

from __future__ import annotations

import hashlib
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, Sequence

from docx import Document
from openai import OpenAI
from pinecone import Pinecone, ServerlessSpec
from pypdf import PdfReader


@dataclass(frozen=True)
class Settings:
    nebius_api_key: str
    pinecone_api_key: str
    nebius_base_url: str = "https://api.tokenfactory.nebius.com/v1/"
    chat_model: str = "Qwen3-235B-A22B-Instruct-2507"
    embedding_model: str = "BAAI/bge-en-icl"
    index_name: str = "claims-rag"
    namespace: str = "claim-clm10001"
    cloud: str = "aws"
    region: str = "us-east-1"
    documents_dir: Path = Path("Documents")
    chunk_size: int = 1200
    chunk_overlap: int = 200
    top_k: int = 5

    @classmethod
    def from_env(cls) -> "Settings":
        return cls(
            nebius_api_key=os.getenv("NEBIUS_API_KEY", ""),
            pinecone_api_key=os.getenv("PINECONE_API_KEY", ""),
            nebius_base_url=os.getenv("NEBIUS_BASE_URL", cls.nebius_base_url),
            chat_model=os.getenv("NEBIUS_CHAT_MODEL", cls.chat_model),
            embedding_model=os.getenv("NEBIUS_EMBEDDING_MODEL", cls.embedding_model),
            index_name=os.getenv("PINECONE_INDEX_NAME", cls.index_name),
            namespace=os.getenv("PINECONE_NAMESPACE", cls.namespace),
            cloud=os.getenv("PINECONE_CLOUD", cls.cloud),
            region=os.getenv("PINECONE_REGION", cls.region),
            documents_dir=Path(os.getenv("DOCUMENTS_DIR", "Documents")),
            chunk_size=int(os.getenv("CHUNK_SIZE", "1200")),
            chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "200")),
            top_k=int(os.getenv("TOP_K", "5")),
        )

    def validate(self) -> None:
        missing = []
        if not self.nebius_api_key:
            missing.append("NEBIUS_API_KEY")
        if not self.pinecone_api_key:
            missing.append("PINECONE_API_KEY")
        if missing:
            raise ValueError(f"Missing environment variables: {', '.join(missing)}")
        if self.chunk_overlap >= self.chunk_size:
            raise ValueError("CHUNK_OVERLAP must be smaller than CHUNK_SIZE")


def read_documents(folder: Path) -> Iterator[tuple[str, int, str]]:
    """Yield (filename, page/section number, text) from supported files."""
    if not folder.exists():
        raise FileNotFoundError(f"Documents directory not found: {folder.resolve()}")
    for path in sorted(folder.iterdir()):
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            for page_number, page in enumerate(PdfReader(path).pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    yield path.name, page_number, text
        elif suffix == ".docx":
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Include table cells; many claim forms are table-based.
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(paragraphs)
            if text.strip():
                yield path.name, 1, text


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Make overlapping chunks, preferring paragraph/sentence boundaries."""
    clean = re.sub(r"[ \t]+", " ", text.replace("\r", ""))
    clean = re.sub(r"\n{3,}", "\n\n", clean).strip()
    if not clean:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(clean):
        end = min(start + size, len(clean))
        if end < len(clean):
            candidates = [clean.rfind("\n", start + size // 2, end), clean.rfind(". ", start + size // 2, end)]
            boundary = max(candidates)
            if boundary > start:
                end = boundary + 1
        chunk = clean[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(clean):
            break
        start = max(end - overlap, start + 1)
    return chunks


class RAGService:
    def __init__(self, settings: Settings):
        settings.validate()
        self.settings = settings
        self.ai = OpenAI(base_url=settings.nebius_base_url, api_key=settings.nebius_api_key)
        self.pc = Pinecone(api_key=settings.pinecone_api_key)

    def _embed(self, texts: Sequence[str]) -> list[list[float]]:
        response = self.ai.embeddings.create(model=self.settings.embedding_model, input=list(texts))
        return [item.embedding for item in sorted(response.data, key=lambda item: item.index)]

    def _ensure_index(self, dimension: int):
        names = [item.name for item in self.pc.list_indexes()]
        if self.settings.index_name not in names:
            self.pc.create_index(
                name=self.settings.index_name,
                dimension=dimension,
                metric="cosine",
                spec=ServerlessSpec(cloud=self.settings.cloud, region=self.settings.region),
            )
            for _ in range(60):
                if self.pc.describe_index(self.settings.index_name).status.get("ready"):
                    break
                time.sleep(1)
        description = self.pc.describe_index(self.settings.index_name)
        if description.dimension != dimension:
            raise ValueError(
                f"Index '{self.settings.index_name}' has dimension {description.dimension}, "
                f"but {self.settings.embedding_model} returned {dimension}. Use a new index name."
            )
        return self.pc.Index(self.settings.index_name)

    def ingest(self, batch_size: int = 32) -> int:
        records = []
        for filename, page, text in read_documents(self.settings.documents_dir):
            for chunk_number, chunk in enumerate(
                chunk_text(text, self.settings.chunk_size, self.settings.chunk_overlap), start=1
            ):
                digest = hashlib.sha256(f"{filename}:{page}:{chunk_number}:{chunk}".encode()).hexdigest()[:32]
                records.append((digest, chunk, {"source": filename, "page": page, "chunk": chunk_number, "text": chunk}))
        if not records:
            raise ValueError("No extractable text found in PDF or DOCX files")

        first_vectors = self._embed([record[1] for record in records[:batch_size]])
        index = self._ensure_index(len(first_vectors[0]))
        # Re-indexing is deterministic and removes stale chunks from this namespace.
        index.delete(delete_all=True, namespace=self.settings.namespace)
        for offset in range(0, len(records), batch_size):
            batch = records[offset : offset + batch_size]
            vectors = first_vectors if offset == 0 else self._embed([record[1] for record in batch])
            index.upsert(
                vectors=[{"id": record[0], "values": vector, "metadata": record[2]} for record, vector in zip(batch, vectors)],
                namespace=self.settings.namespace,
            )
        return len(records)

    def ask(self, question: str) -> tuple[str, list[dict]]:
        vector = self._embed([question])[0]
        index = self.pc.Index(self.settings.index_name)
        result = index.query(
            vector=vector,
            top_k=self.settings.top_k,
            include_metadata=True,
            namespace=self.settings.namespace,
        )
        sources = []
        context_parts = []
        for number, match in enumerate(result.matches, start=1):
            metadata = dict(match.metadata or {})
            source = {"number": number, "source": metadata.get("source", "Unknown"), "page": metadata.get("page"), "score": float(match.score)}
            sources.append(source)
            context_parts.append(f"[{number}] {source['source']}, page/section {source['page']}\n{metadata.get('text', '')}")
        if not context_parts:
            return "I could not find relevant information in the indexed documents.", []
        system = (
            "You answer questions only from the supplied claim-document context. "
            "If the context is insufficient, say so. Cite supporting passages inline as [1], [2], etc. "
            "Do not invent claim facts, policy terms, amounts, dates, or citations."
        )
        prompt = f"Context:\n\n{chr(10).join(context_parts)}\n\nQuestion: {question}"
        response = self.ai.chat.completions.create(
            model=self.settings.chat_model,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            temperature=0.1,
        )
        return response.choices[0].message.content or "No answer returned.", sources

